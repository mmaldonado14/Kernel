[markdown]
# Connected to Python 3.13.1

# %%
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
import feedparser

# Fuentes RSS relevantes (puedes agregar más)
RSS_FEEDS = [
    "https://news.google.com/rss/search?q=memorándum+de+entendimiento+Panamá+Estados+Unidos",
    "https://www.prensa.com/rss/ultimas-noticias/",
    "https://www.laestrella.com.pa/rss",
    "https://feeds.bbci.co.uk/mundo/rss.xml",
]

# Función para obtener país y medio de comunicación
def get_source_info(entry):
    if 'source' in entry:
        medio = entry['source']['title']
    elif 'title_detail' in entry:
        medio = entry['title_detail']['base']
    else:
        medio = entry.get('link', '').split('/')[2] if 'link' in entry else 'Desconocido'
    # País por heurística simple
    if '.pa' in medio or 'panama' in medio:
        pais = 'Panamá'
    elif '.us' in medio or 'unitedstates' in medio or 'america' in medio:
        pais = 'Estados Unidos'
    elif '.mx' in medio or 'mexico' in medio:
        pais = 'México'
    elif '.es' in medio or 'spain' in medio:
        pais = 'España'
    elif '.uk' in medio or 'bbc' in medio:
        pais = 'Reino Unido'
    else:
        pais = 'Internacional'
    return medio, pais

# Recopilar noticias
noticias = []
for feed_url in RSS_FEEDS:
    feed = feedparser.parse(feed_url)
    for entry in feed.entries:
        titulo = entry.title
        link = entry.link
        medio, pais = get_source_info(entry)
        # Filtrar solo noticias relevantes
        if 'memorándum' in titulo.lower() or 'memorandum' in titulo.lower():
            noticias.append({
                'titulo': titulo,
                'medio': medio,
                'pais': pais,
                'link': link
            })print ("Hola Mundo")

# Crear documento Word
doc = Document()
doc.add_heading('Reporte de Noticias: Memorándum de Entendimiento Panamá - Estados Unidos', 0)
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Título de la Noticia'
hdr_cells[1].text = 'Medio de Comunicación'
hdr_cells[2].text = 'País'

for noticia in noticias:
    row_cells = table.add_row().cells
    row_cells[0].text = noticia['titulo']
    row_cells[1].text = noticia['medio']
    row_cells[2].text = noticia['pais']

doc.add_paragraph(f'Total de noticias encontradas: {len(noticias)}')
doc.save('Reporte_Memorandum_Panama_EEUU.docx')
print("Reporte generado: Reporte_Memorandum_Panama_EEUU.docx")


